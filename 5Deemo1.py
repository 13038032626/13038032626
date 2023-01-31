import docx
from docx import Document
import random

file = Document("D:\python\课程设计Part2\选择题.docx")
timu_list = []
daan_list = []
for para in file.paragraphs:
    if para.text not in ['A', 'B', 'C', 'D']:
        timu_list.append(para.text)
    else:
        daan_list.append(para.text)

#print(timu_list)
#print('::::::::::')
#print(daan_list)
try:
    file1 = Document(r"D:\python\课程设计Part2\5题\id和密码.docx")
    a = input("输入你的id")
    for para in file1.paragraphs:
        if a in para.text:
            b = input("输入密码")  # 模块化下像用用户名有点困难
            if b in para.text:
                print("登陆成功")
                break
    else:
        raise TypeError


    def cheshi():
        chengji = 0
        doc = docx.Document(r"D:\python\课程设计Part2\5题\test.docx")
        # paragraph=doc.add_paragraph()
        li = [x for x in range(10)]
        random.shuffle(li)
        new_li = li[0:5]
        print("考生姓名：", a)
        print("答题开始")
        print("{:*^100}".format('————————————————'))
        for i in range(5):
            answer = input('{}\n{}'.format(timu_list[new_li[i] * 2], timu_list[new_li[i] * 2 + 1]))
            paragraph1 = doc.add_paragraph(timu_list[new_li[i] * 2])
            paragraph2 = doc.add_paragraph(timu_list[new_li[i] * 2 + 1])
            doc.save(r"D:\python\课程设计Part2\5题\{}.docx".format(a))
            if answer == daan_list[new_li[i]]:
                chengji += 1
        return chengji


    print('成绩是', cheshi())

except TypeError:
    print("登陆失败")
